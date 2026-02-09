#!/usr/bin/env python3
"""Generate Rossi Mission Project Product Document in Word format.
V2 — Palette-informed rewrite with marketplace research evidence."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# -- Style Setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


def add_styled_paragraph(text, bold=False, italic=False, size=None, color=None,
                         alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_evidence(text):
    """Add an evidence marker in italic gray."""
    p = doc.add_paragraph()
    run = p.add_run(f'[Evidence: {text}]')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x99)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table_from_data(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ================================================================
# COVER PAGE
# ================================================================
for _ in range(6):
    doc.add_paragraph()

add_styled_paragraph('ROSSI MISSION PROJECT', bold=True, size=28,
                      color=(0x1A, 0x1A, 0x2E),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_styled_paragraph('Product Strategy & Discovery Plan', bold=True, size=18,
                      color=(0x4A, 0x4A, 0x6A),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_styled_paragraph('A product-craft approach to validating and scaling\n'
                      'a cultural platform for graffiti artists',
                      italic=True, size=13,
                      color=(0x6A, 0x6A, 0x8A),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_styled_paragraph('Prepared for Rossi Leadership Team', size=11,
                      color=(0x4A, 0x4A, 0x4A),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_styled_paragraph('February 2026  |  V2 — Palette System Validated', size=11,
                      color=(0x4A, 0x4A, 0x4A),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_styled_paragraph('Product Management Practice', size=11,
                      color=(0x4A, 0x4A, 0x4A),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONFIDENTIAL')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)

doc.add_page_break()

# ================================================================
# TABLE OF CONTENTS
# ================================================================
doc.add_heading('Contents', level=1)
toc_items = [
    '1. The Customer Problem',
    '2. What We Know (Evidence)',
    '3. The Market: Who Rossi Competes With',
    '4. What We Believe (Hypotheses)',
    '5. What We Need to Learn (Discovery Plan)',
    '6. The Product: What Rossi Actually Sells',
    '7. Product Strategy (Year 1-3)',
    '8. Product Risks and How We Reduce Them',
    '9. The Team',
    '10. Success Metrics',
    '11. 90-Day Discovery Roadmap',
    '12. Recommendation to Leadership',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ================================================================
# SECTION 1: THE CUSTOMER PROBLEM
# ================================================================
doc.add_heading('1. The Customer Problem', level=1)

doc.add_paragraph(
    'Before we talk about Rossi\'s product, we need to talk about the customer problem '
    'Rossi exists to solve. Everything in this document follows from this.'
)

doc.add_heading('The Artist\'s Problem', level=2)
doc.add_paragraph(
    'Graffiti and street artists face a career ceiling that other artists do not. '
    'A painter can go from art school to gallery representation to museum shows in a '
    'linear path. A graffiti artist has no such path. The traditional art world treats '
    'street art as illegitimate. Streetwear brands treat artists as vendors, paying flat '
    'fees for designs and keeping the ongoing revenue. Online platforms like Etsy and Artsy '
    'offer distribution but no curation, no career development, and no community.'
)

doc.add_paragraph(
    'Our marketplace research identified 50+ graffiti artists with active merchandise '
    'operations. The pattern is consistent: established artists like Cope2, Lady Pink, '
    'and Saber run their own D2C shops (Cope2-Merch.com, LadyPinkNYC.com, SaberOne.com) '
    'because no institution serves them. They had to build their own infrastructure. '
    'Emerging artists — the ones without 200K+ Instagram followers — have no path at all.'
)
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — 50+ artists mapped across online + physical channels')

doc.add_paragraph(
    'The result: talented artists with signature styles, cultural credibility, and community '
    'followings who cannot convert that into sustainable creative careers.'
)

doc.add_heading('The Collector\'s Problem', level=2)
doc.add_paragraph(
    'Art collectors interested in street art have no trusted intermediary. Traditional '
    'galleries charge 40-50% commission and don\'t carry graffiti art. Online platforms '
    'like Etsy have 10,000+ graffiti-related listings but no quality filter and massive '
    'unauthorized reproduction problems. StockX provides authentication for established '
    'artists (KAWS, Futura 2000) but nothing for emerging talent. There is no curated '
    'discovery channel for this market segment.'
)
add_evidence('Marketplace research: Etsy 10K+ listings, 6.5% fees, no curation; StockX 9.5% seller fee, '
             'authentication only for established artists')

doc.add_heading('The Community\'s Problem', level=2)
doc.add_paragraph(
    'The Mission District has 40+ years of graffiti culture. It produced the Mission School '
    'movement. It is one of the most culturally significant neighborhoods in American street '
    'art. But the artists who create that culture see almost no economic benefit from it. '
    'The community wants an institution that is OF the neighborhood, not extracting FROM it.'
)

doc.add_heading('The Problem Statement', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'There is no institution that systematically discovers, develops, and launches '
    'graffiti artist careers while keeping artists as economic partners rather than vendors.'
)
run.bold = True
run.italic = True
doc.add_paragraph(
    'This is the customer problem Rossi exists to solve. Everything that follows — '
    'every product decision, every revenue stream, every partnership — must be evaluated '
    'against this problem. If it helps solve it, we do it. If it doesn\'t, we don\'t.'
)

doc.add_page_break()

# ================================================================
# SECTION 2: WHAT WE KNOW (EVIDENCE)
# ================================================================
doc.add_heading('2. What We Know (Evidence)', level=1)

doc.add_paragraph(
    'Product decisions should be grounded in evidence, not opinions. Here is what we '
    'actually know, supported by data, about the market Rossi operates in.'
)

doc.add_heading('The Market Is Real', level=2)
add_bullet('Global street art market: $6.8B, growing at 8.2% CAGR through 2030.')
add_bullet('Artist brand monetization (music industry analog): $3.5B in expanded rights.')
add_bullet('Arts education in California: $2B+ market.')
add_bullet('Valencia Street retail vacancy: 2-4% vs 6-8% citywide. The location is an asset.')
add_bullet('50+ graffiti artists identified with active merchandise operations across 7+ platforms.')
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md; research/argy/task-1-global-markets.md')

doc.add_heading('The Revenue Models Are Proven', level=2)
doc.add_paragraph(
    'Our marketplace research mapped six distinct revenue models used by working graffiti artists:'
)
add_table_from_data(
    ['Model', 'Examples', 'Artist Take', 'Rossi Advantage'],
    [
        ['Direct-to-Consumer', 'Cope2-Merch.com, SaberOne.com', '70-90%', 'Rossi adds curation + community'],
        ['Gallery Representation', 'Retna, Mear One, Os Gemeos', '50-60%', 'Rossi matches at 50/50'],
        ['Print Editions', 'DirtyPilot, Edition Copenhagen', '50%', 'Rossi adds career development'],
        ['Marketplace (Etsy, StockX)', 'Various', '80-93.5%', 'Rossi adds authentication + brand'],
        ['Brand Collaborations', 'KAWS x Dior, Futura x Kenzo', '5-15% royalty', 'Rossi brokers at 70/30 artist-favorable'],
        ['Artist Collectives', 'Alberta Street (70/30), Side Rail (100%)', '70-100%', 'Rossi adds scale + online platform'],
    ]
)
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — Revenue Models Observed section')

doc.add_heading('The Model Works at Scale', level=2)
doc.add_paragraph(
    'Creative Growth Art Center in Oakland operates Rossi\'s exact model at scale: '
    'nonprofit gallery, 50/50 artist split, 140 artists, $1.06M in annual art sales, '
    '50-year track record, SFMOMA partnership.'
)
add_table_from_data(
    ['Metric', 'Creative Growth (FY 2024)', 'Rossi (Year 1 Projection)'],
    [
        ['Total Revenue', '$3.26M', '$745K'],
        ['Art Sales', '$1.06M (33%)', '$350K (47%)'],
        ['Grants', '$1.11M (34%)', '$110K (15%)'],
        ['Programs', '$1.03M (32%)', '$285K (38%)'],
        ['Artists Served', '140', '80+'],
        ['Gross Revenue / Artist', '$7,593', '$4,375'],
    ]
)
add_evidence('Creative Growth Form 990 FY 2024, EIN 23-7319028, ProPublica Nonprofit Explorer')

doc.add_heading('Creative Growth: Historical Trend (2020-2024)', level=3)
doc.add_paragraph(
    'Five years of Form 990 data reveal both the sustainability and the volatility '
    'inherent in this model:'
)
add_table_from_data(
    ['Year', 'Revenue', 'Art Sales', 'Grants', 'Net Income'],
    [
        ['2020', '$2.83M', '$1.00M (35%)', '$1.09M (39%)', '+$356K'],
        ['2021', '$2.81M', '$806K (29%)', '$1.10M (39%)', '+$51K'],
        ['2022', '$3.40M', '$954K (28%)', '$1.39M (41%)', '+$762K'],
        ['2023', '$2.41M', '$499K (21%)', '$732K (30%)', '-$643K'],
        ['2024', '$3.26M', '$1.06M (33%)', '$1.11M (34%)', '-$741K'],
    ]
)
doc.add_paragraph(
    'Two critical patterns emerge. First, art sales fluctuate between 21% and 35% of '
    'revenue year to year — a nearly 2x swing ($499K to $1.06M). Retail-dependent projections '
    'must account for this volatility. Second, grants are the stabilizer: consistently 30-41% '
    'of revenue across all five years. Creative Growth survives art sales volatility because '
    'its grant base is diversified and deep. Rossi must build the same resilience.'
)
doc.add_paragraph(
    'Also notable: Creative Growth spends 57% of expenses on personnel ($2.28M), compared '
    'to Rossi\'s projected 28% ($241K). As Rossi scales — adding a grant writer, workshop '
    'coordinator, and content team — personnel costs will grow toward Creative Growth\'s '
    'proportion. Year 2-3 expense projections should anticipate this.'
)
add_evidence('Creative Growth Form 990s 2020-2024, ProPublica Nonprofit Explorer')

p = doc.add_paragraph()
run = p.add_run(
    'The positioning statement: "Creative Growth Art Center in Oakland has operated a nearly '
    'identical model for 50 years: nonprofit gallery, 50/50 artist split, 140 artists, $1M+ '
    'in annual art sales. Rossi is the graffiti and street art version of Creative Growth, '
    'serving an underserved artist community in SF\'s Mission District." This is infinitely '
    'more credible than comparisons to Supreme or Stussy.'
)
run.italic = True

doc.add_heading('Two Artist Collectives Validate the Partner Model', level=2)
doc.add_paragraph(
    'Our marketplace research identified two existing cooperatives that validate '
    'Rossi\'s partner expansion targets:'
)
add_bullet('Alberta Street Gallery, Portland (1829 NE Alberta St) — 30 member artists, '
           '70/30 artist/gallery split. Proves cooperative model works in a target city.', bold_prefix='Portland: ')
add_bullet('Side Rail Collective, Seattle (5511 1/2 Airport Way S, Georgetown) — '
           'No-commission model. Artists keep 100%, pay membership fees ($50-200/month). '
           'Proves artist-first economics attract talent.', bold_prefix='Seattle: ')
add_bullet('Gray Loft Gallery, Oakland (489 25th Street, Jingletown) + Oakland Art Murmur — '
           'Active artist collective scene, First Friday gallery walks. Validates Oakland as Year 2 target.', bold_prefix='Oakland: ')
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — Physical Locations section')

doc.add_heading('Rossi Already Has Product in Market', level=2)
add_bullet('80+ artists on the roster (the supply side exists).')
add_bullet('251 SKUs on a physical shelf at 791 Valencia Street.')
add_bullet('An operational Shopify store (digital channel exists, even if underdeveloped).')
add_bullet('A location in one of SF\'s healthiest retail corridors.')

doc.add_heading('The Critical Gap: We Don\'t Know Our Own Revenue', level=2)
p = doc.add_paragraph()
run = p.add_run('We do not know Rossi\'s current revenue. ')
run.bold = True
p.add_run(
    'The business plan projects $780K Year 1, but no one has pulled the POS data. '
    'The modeled estimate is $245K-$390K annually. The difference between $250K and $400K '
    'determines whether Year 1 targets are ambitious or unrealistic.'
)
add_evidence('UNDERWRITER_AUDIT_FINAL.md — Condition 1 assessment; condition-1-trailing-actuals.md')

doc.add_page_break()

# ================================================================
# SECTION 3: THE MARKET — COMPETITIVE LANDSCAPE
# ================================================================
doc.add_heading('3. The Market: Who Rossi Competes With', level=1)

doc.add_paragraph(
    'Our marketplace research mapped the complete ecosystem of graffiti artist merchandise '
    'sales — online platforms, physical galleries, and direct sales channels. Rossi does not '
    'compete with traditional galleries or generic online platforms. It competes with the '
    'fragmented system that artists have built for themselves.'
)

doc.add_heading('Direct Competitors', level=2)
add_table_from_data(
    ['Competitor', 'Model', 'Strength', 'Weakness', 'Rossi Advantage'],
    [
        ['GraffitiStreet.com', 'Online gallery, 100+ artists', 'Global shipping, established',
         '40-50% commission', 'Better economics (50/50)'],
        ['DirtyPilot.com', 'Curated prints, certificates', 'Authentication, quality',
         'Online only, no physical', 'Hybrid: physical + online'],
        ['Alberta Street Gallery', 'Cooperative, 30 members', '70/30 split, community',
         'Local only, no scale', 'National network + online'],
        ['Side Rail Collective', 'No-commission, membership', 'Artists keep 100%',
         'Local only, no curation', 'Curation + career development'],
        ['Etsy', 'Marketplace, 10K+ listings', 'Massive reach, low fees (6.5%)',
         'No curation, fake reproductions', 'Authentication + brand'],
    ]
)
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — Competitive Landscape section')

doc.add_heading('How Top Artists Actually Sell Today', level=2)
doc.add_paragraph(
    'The most successful graffiti artists use 3-5 channels simultaneously. This validates '
    'Rossi\'s multi-channel approach — but also reveals the gap: only established artists '
    'with large followings can run this themselves. Emerging artists cannot.'
)
add_table_from_data(
    ['Artist', 'Followers', 'Channels Used', 'Price Range'],
    [
        ['Shepard Fairey / OBEY', '2.1M', 'Own brand + retailers + gallery', '$30-$100K'],
        ['KAWS', '4.8M', 'DDT Store drops + Dior/Uniqlo collabs + galleries', '$200-$15M'],
        ['Cope2', '400K', 'Cope2-Merch.com + DirtyPilot + gallery shows', '$200-$25K'],
        ['Lady Pink', '200K', 'LadyPinkNYC.com + Wynwood Gallery + collabs', '$50-$50K'],
        ['Saber', '250K', 'SaberOne.com + Signari Gallery', '$500-$100K'],
    ]
)
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — Tier 1 & Tier 2 artist profiles')

doc.add_paragraph(
    'The insight: artists with 100K+ followers can build their own infrastructure. '
    'Artists with 1K-50K followers — the emerging tier — cannot. That is Rossi\'s market.'
)

doc.add_heading('Rossi\'s Competitive Position', level=2)
add_table_from_data(
    ['Dimension', 'Rossi', 'Traditional Gallery', 'Online Platform', 'Artist Collective'],
    [
        ['Artist Economics', '50/50 split', '40-60% to gallery', '6.5-20% platform fee', '70-100% to artist'],
        ['Career Development', 'Full pipeline', 'None', 'None', 'Peer only'],
        ['Cultural Credibility', 'Very High (Mission District)', 'High (curator-driven)', 'Medium', 'Medium'],
        ['Online Presence', 'Low (fixable)', 'Low', 'Very High', 'Low'],
        ['Scarcity Discipline', 'High (250 SKU cap)', 'None', 'None', 'None'],
        ['Authentication', 'Yes', 'Yes', 'Only StockX', 'No'],
    ]
)
add_evidence('BCG_PRODUCT_ASSESSMENT.md Section 3.3; GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md')

doc.add_paragraph(
    'Rossi\'s gap is digital presence. Every other dimension is strong or differentiated. '
    'The online strategy is the single biggest product vulnerability — and the highest-ROI fix.'
)

doc.add_page_break()

# ================================================================
# SECTION 4: WHAT WE BELIEVE (HYPOTHESES)
# ================================================================
doc.add_heading('4. What We Believe (Hypotheses)', level=1)

doc.add_paragraph(
    'In product work, we distinguish between what we know (evidence) and what we believe '
    '(hypotheses). Hypotheses are essential. But they need to be tested before we build '
    'products around them.'
)

add_table_from_data(
    ['Hypothesis', 'Risk', 'Current Evidence', 'Test'],
    [
        ['Workshops discover talent (10% conversion)',
         'HIGH', 'Zero workshops conducted', 'Run 1 pilot, track 90 days'],
        ['Content drives art sales (GX1000 model)',
         'HIGH', 'No attribution data', 'Tag links, track funnel 60 days'],
        ['Scarcity works for emerging artists',
         'MEDIUM', 'Supreme model — but Rossi is not Supreme', 'Monitor sell-through 6 months'],
        ['Artists stay because of 50/50 split',
         'MEDIUM', 'Alberta St (70/30) + Side Rail (100%) retain artists', 'Survey 20 artists'],
        ['Partner galleries will join network',
         'MEDIUM', 'Alberta St + Side Rail prove cooperatives exist', 'Pitch 3 galleries, get 1 LOI'],
        ['$1,500-3,000 corporate workshops sell',
         'MEDIUM', 'SFMOMA charges $2K-$4K (market validated)', 'Book 1 workshop within 90 days'],
        ['Online can reach $55K Year 1',
         'LOW', 'Cope2, Lady Pink run D2C shops successfully', 'Optimize store, run drops, measure 90 days'],
        ['Brand collabs generate $70K+',
         'MEDIUM', 'KAWS x Uniqlo, Futura x Kenzo prove market', 'Pitch 2 brands, close 1'],
    ]
)
add_evidence('Palette Anky validation framework; marketplace research comparables')

doc.add_paragraph(
    'Every hypothesis above can be tested for under $2,000 and within 90 days. '
    'There is no reason to build a 7-stream operation on untested assumptions.'
)

doc.add_page_break()

# ================================================================
# SECTION 5: DISCOVERY PLAN
# ================================================================
doc.add_heading('5. What We Need to Learn (Discovery Plan)', level=1)

doc.add_paragraph(
    'Product discovery is how we reduce risk. We identify the biggest unknowns, design '
    'the cheapest experiments to resolve them, and make investment decisions based on evidence.'
)

doc.add_heading('Sprint 1: Do We Have a Business? (Days 1-14)', level=2)
add_bullet('Pull 12 months of POS data from Square/Shopify.', bold_prefix='Action: ')
add_bullet('We learn our actual revenue, transaction volume, top artists, and seasonality.', bold_prefix='Outcome: ')
add_bullet('2 days of work. Zero cost.', bold_prefix='Cost: ')
add_bullet('If revenue is $300K+, the plan is credible. If below $200K, major revision.', bold_prefix='Gate: ')

doc.add_heading('Sprint 2: Does the Pipeline Work? (Days 14-45)', level=2)
add_bullet('Interview 3-5 existing roster artists. Document before/after.', bold_prefix='Action 1: ')
add_bullet('Run 1 pilot workshop (12-15 participants, 4 hours, $75-125 per person).', bold_prefix='Action 2: ')
add_bullet('Survey 20 artists on satisfaction, retention intent, and improvement ideas.', bold_prefix='Action 3: ')
add_bullet('Evidence that Rossi has already changed artist careers + workshop attracts participants.', bold_prefix='Outcome: ')
add_bullet('~$500.', bold_prefix='Cost: ')

doc.add_heading('Sprint 3: Can We Sell Online? (Days 14-45, parallel)', level=2)
doc.add_paragraph(
    'Our marketplace research shows that every successful graffiti artist uses a hybrid model — '
    'direct website, Instagram shop, gallery representation, and print editions. Rossi must build '
    'this infrastructure for its artists.'
)
add_bullet('Hire freelance photographer. Shoot top 50 products.', bold_prefix='Action 1: ')
add_bullet('Optimize Shopify (collections, SEO, email capture, print-on-demand via Printful).', bold_prefix='Action 2: ')
add_bullet('Run first "Artist Drop" (featured artist, 3-5 pieces, Instagram + Shopify).', bold_prefix='Action 3: ')
add_bullet('Measurable online sales data from a real product launch.', bold_prefix='Outcome: ')
add_bullet('~$6,000.', bold_prefix='Cost: ')
add_bullet('If first artist drop generates $500+ in 72 hours, the channel works.', bold_prefix='Gate: ')
add_evidence('Marketplace research: Instagram drops are standard for artists with 100K+ followers')

doc.add_heading('Sprint 4: Will Partners Join? (Days 45-90)', level=2)
doc.add_paragraph(
    'We now have specific targets. The marketplace research identified real galleries in '
    'every expansion city:'
)
add_bullet('Oakland — Gray Loft Gallery (489 25th St, Jingletown), Oakland Art Murmur network')
add_bullet('Portland — Alberta Street Gallery (1829 NE Alberta St, 70/30 split, 30 members)')
add_bullet('Seattle — Side Rail Collective (5511 1/2 Airport Way S, no-commission model)')
add_bullet('LA — 01 Gallery (Melrose), 33 1/3 Gallery (Silverlake)')
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — Physical Locations section')
add_bullet('Have conversations with 3 galleries. Aim for 1 letter of intent.', bold_prefix='Action: ')
add_bullet('Zero cost.', bold_prefix='Cost: ')

doc.add_page_break()

# ================================================================
# SECTION 6: THE PRODUCT
# ================================================================
doc.add_heading('6. The Product: What Rossi Actually Sells', level=1)

doc.add_paragraph(
    'After reviewing the complete business plan, marketplace research, financial models, '
    'and competitive analysis, our assessment is that Rossi sells three products under a '
    'unified cultural platform brand.'
)

doc.add_heading('Product 1: Curated Art (Gallery + Online)', level=2)
doc.add_paragraph(
    'Physical and digital retail of original art, limited prints, and artist merchandise, '
    'curated by the artist collective and capped at 250 SKUs per location. This is the '
    'core revenue driver (47% of Year 1 revenue).'
)
doc.add_paragraph(
    'The competitive advantage: Rossi\'s 50/50 split is better than traditional galleries '
    '(40-60% to gallery) while offering curation and authentication that platforms like '
    'Etsy (6.5% fee, no quality control) cannot match.'
)
add_evidence('Marketplace research commission comparison; Creative Growth Form 990')
p = doc.add_paragraph()
run = p.add_run('Product-market fit status: Validated at current scale.')
run.bold = True

doc.add_heading('The Artist Revenue Distribution (Pareto Reality)', level=3)
doc.add_paragraph(
    'Creative Growth\'s 50-year track record confirms what every gallery knows: artist '
    'revenue follows a Pareto distribution. Top 10% of artists generate 60-70% of sales. '
    'For Rossi\'s 80 artists at $350K retail, the math is honest:'
)
add_table_from_data(
    ['Tier', 'Artists', 'Gross Sales', 'Per Artist (Net)', 'Role in Pipeline'],
    [
        ['Top 10%', '8 artists', '$210-245K (60-70%)', '$13K-15K/year', 'Revenue anchors — fund the operation'],
        ['Middle 30%', '24 artists', '$70-88K (20-25%)', '$1,450-1,830/year', 'Growth tier — pipeline working'],
        ['Emerging 60%', '48 artists', '$18-70K (5-20%)', '$185-730/year', 'Development tier — building toward mid'],
    ]
)
doc.add_paragraph(
    'This is not a bug. This is how art markets work. The pipeline\'s job is to move '
    'artists UP the distribution over time: emerging artists build skills through workshops, '
    'gain visibility through content and events, and graduate into the middle and top tiers. '
    'The metric that matters is not whether every artist earns equally — it is whether artists '
    'are moving up the curve year over year.'
)
add_evidence('VALIDATED_COMPARABLES_RESEARCH.md — Finding 6: Pareto Distribution')

doc.add_heading('Product 2: Experiences (Workshops + Events + Residencies)', level=2)
doc.add_paragraph(
    'Skill-building workshops, corporate team-building experiences, gallery events, and '
    'sponsored artist residencies. This serves two functions: revenue generation and '
    'talent discovery (the pipeline). Pricing validated against market: SFMOMA charges '
    '$2,000-$4,000 for corporate workshops; Museum of Craft & Design charges $1,500-$2,500.'
)
add_evidence('UNDERWRITER_READY_REPORT.md Section 6 — workshop pricing benchmarks')
p = doc.add_paragraph()
run = p.add_run('Product-market fit status: Unvalidated.')
run.bold = True
p.add_run(' Zero workshops have been conducted.')

doc.add_heading('Product 3: Brand Partnerships (Collaborations + Licensing)', level=2)
doc.add_paragraph(
    'Curated brand collaborations between Rossi artists and external partners. Capped at '
    '3-5 per year to protect scarcity. Artist veto power preserved. 70% of revenue to artist.'
)
doc.add_paragraph(
    'The market is proven at the top tier: KAWS x Dior, KAWS x Uniqlo, Futura 2000 x Kenzo, '
    'Stash x Nike, Lady Pink x Weatherman Umbrellas. Licensing fees range from $10K-$500K '
    'with 5-15% royalties. Rossi targets the $5K-$30K collaboration range for emerging artists '
    '— a tier currently unserved by any intermediary.'
)
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — Brand Collaborations section')
p = doc.add_paragraph()
run = p.add_run('Product-market fit status: Unvalidated.')
run.bold = True
p.add_run(' No collaborations completed. But the market is proven by KAWS, Futura, and dozens of others.')

doc.add_heading('The Brand Architecture', level=2)
add_bullet('The master brand and cultural platform identity.', bold_prefix='Rossi Mission Project: ')
add_bullet('The physical and online retail experience.', bold_prefix='Rossi Gallery: ')
add_bullet('Artist development programs (workshops, residencies, mentorship).', bold_prefix='Rossi Studio: ')
add_bullet('Limited-edition product drops (prints, merch, collaborations).', bold_prefix='Rossi Editions: ')
add_bullet('The content and documentation engine.', bold_prefix='Rossi Stories: ')
doc.add_paragraph(
    'This architecture lets each product line speak to its natural audience while '
    'maintaining coherence. Grant funders engage with Rossi Studio. Collectors engage '
    'with Rossi Gallery. Streetwear consumers engage with Rossi Editions.'
)

doc.add_page_break()

# ================================================================
# SECTION 7: PRODUCT STRATEGY
# ================================================================
doc.add_heading('7. Product Strategy', level=1)

doc.add_heading('Year 1: Validate the Core', level=2)
doc.add_paragraph(
    'Year 1 is not about growth. Year 1 is about learning. Validate that the three '
    'products work at the SF flagship before investing in expansion.'
)

add_table_from_data(
    ['Product', 'Year 1 Goal', 'Target Revenue', 'Investment'],
    [
        ['Curated Art (Gallery)', 'Establish baseline, grow 15-25%', '$350K', 'Existing operations'],
        ['Curated Art (Online)', 'Launch proper e-commerce', '$55K', '$6K'],
        ['Workshops', 'Validate pipeline hypothesis', '$80K (6+ workshops)', '$2K pilot'],
        ['Events', 'Sustain community engagement', '$60K (12 events)', 'Existing'],
        ['Collaborations', 'Build pipeline, close 2-3 deals', '$70K', '$0 (brand deck)'],
        ['Residencies', 'Prove model with 2-3 pilots', '$20K', '$0 (outreach)'],
        ['Grants', 'Submit applications', '$110K', '$0-40K (grant writer)'],
    ]
)

doc.add_heading('Grant Revenue: Expected Value Math', level=3)
doc.add_paragraph(
    'The $110K grant target requires aggressive application volume. First-time applicant '
    'success rates range from 25-40% (public grants) to 20-30% (private foundations). '
    'Based on validated research of the SF Bay Area grant landscape:'
)
add_table_from_data(
    ['Source', 'Applied', 'Probability', 'Expected Value'],
    [
        ['SFAC Individual (2 artists)', '$50K', '50%', '$25K'],
        ['SFAC Organizational', '$25K', '40%', '$10K'],
        ['CAC Organizational', '$15K', '35%', '$5K'],
        ['Zellerbach Foundation', '$25K', '30%', '$7.5K'],
        ['Haas Fund', '$50K', '25%', '$12.5K'],
        ['SF Foundation', '$50K', '20%', '$10K'],
        ['TOTAL', '$215K applied', '', '$70K expected'],
    ]
)
doc.add_paragraph(
    'To net $110K in grants, Rossi must apply for $300K+ and achieve above-average '
    'success rates. This requires either a dedicated grant writer (0.5 FTE, ~$40K) or '
    'a fiscal sponsor with grant-writing capacity. The $50K grant target in the original '
    'plan was too low; $110K is achievable but demands institutional effort.'
)
add_evidence('VALIDATED_COMPARABLES_RESEARCH.md — Finding 5: Grant Landscape Reality Check')

doc.add_heading('Year 2: Scale What Works', level=2)
doc.add_paragraph(
    'Year 2 decisions are made based on Year 1 evidence. We only scale validated programs.'
)
add_bullet('If workshops validated: expand to 24/year, add Oakland partner.')
add_bullet('If online validated: invest in content production.')
add_bullet('If partner model validated: onboard Oakland (Gray Loft or Art Murmur network).')
add_bullet('If NOT validated: pivot or kill the program before investing further.')
add_evidence('Palette ONE-WAY DOOR framework — expansion is reversible only at high cost')

doc.add_heading('Year 3: Expand Deliberately', level=2)
doc.add_paragraph(
    'Year 3 expansion is contingent on Year 2 evidence. If Oakland works, add LA and one '
    'other city. If not, stay focused until the model is right.'
)

doc.add_heading('Revenue Projections (Discovery-Informed)', level=2)
add_table_from_data(
    ['Year', 'Revenue', 'Growth Rate', 'Locations', 'Milestone'],
    [
        ['Year 1', '$745K', 'Baseline', 'SF only', 'All 3 products validated'],
        ['Year 2', '$1.0-1.1M', '35-45%', 'SF + Oakland', 'Partner model proven'],
        ['Year 3', '$1.4-1.6M', '35-45%', 'SF + Oakland + LA + 1', 'National expansion ready'],
    ]
)
doc.add_paragraph(
    'Note: The original business plan projected $1.55M Year 2 (+99%) and $2.7M Year 3 (+74%). '
    'These are venture-scale growth rates on nonprofit capital. Our projections are 35-45% annually — '
    'still exceptional, but achievable without requiring every variable to go perfectly.'
)
add_evidence('UNDERWRITER_AUDIT_FINAL.md — Year 2-3 growth rates flagged as not credible; '
             'BCG_PRODUCT_ASSESSMENT.md Section 8')

doc.add_page_break()

# ================================================================
# SECTION 8: PRODUCT RISKS
# ================================================================
doc.add_heading('8. Product Risks and How We Reduce Them', level=1)

add_table_from_data(
    ['Risk', 'Probability', 'Impact', 'Early Detection'],
    [
        ['Brand identity confusion', 'High', 'High',
         'Leadership alignment session Week 1. Decide: platform identity.'],
        ['Workshop pipeline fails', 'Medium', 'High',
         'Pilot workshop Month 2-3. If <8 attend, redesign.'],
        ['Content doesn\'t drive sales', 'Medium', 'Medium',
         'Tag all content links Day 1. If zero attribution Month 6, pivot.'],
        ['Online underperforms', 'Low', 'Medium',
         'First artist drop Month 1. If <$500 in 72hrs, adjust.'],
        ['Partner model is net-negative', 'High', 'Medium',
         'Don\'t onboard partners until Year 2. Restructure economics first.'],
        ['Top artists leave', 'Medium', 'High',
         'Artist satisfaction survey Month 1. Address issues early.'],
        ['Month 17 cash crisis ($10K)', 'Medium', 'Critical',
         'Staged funding. Tranche 2 at Month 6 as checkpoint.'],
        ['Grant revenue falls short', 'Medium', 'High',
         'Submit 6-8 applications (not 3-4). Expected value math.'],
        ['Personnel costs escalate', 'High', 'Medium',
         'Creative Growth spends 57% on staff vs Rossi 28%. As grant writer, workshop '
         'coordinator, and content roles are added, personnel will approach 45-55% of expenses. '
         'Budget Year 2-3 accordingly.'],
        ['Rent burden ($120K/year)', 'Medium', 'Medium',
         'Verify actual lease terms. At $10K/month, rent is 16% of projected revenue — '
         'a structural disadvantage vs Oakland-based Creative Growth. Negotiate or plan '
         'for lower-cost Year 2-3 expansion spaces.'],
        ['Art sales volatility', 'High', 'Medium',
         'Creative Growth art sales swung from $499K to $1.06M in one year (2x). '
         'Build 3-month operating reserve. Do not assume steady-state retail.'],
    ]
)
add_evidence('UNDERWRITER_AUDIT_FINAL.md — cash flow danger zones; '
             'BCG_PRODUCT_ASSESSMENT.md Section 10 risk matrix; '
             'VALIDATED_COMPARABLES_RESEARCH.md — Creative Growth historical data')

doc.add_page_break()

# ================================================================
# SECTION 9: THE TEAM
# ================================================================
doc.add_heading('9. The Team', level=1)

doc.add_paragraph(
    'Product outcomes are delivered by empowered teams, not by plans. The question is '
    'whether the team has the capability, autonomy, and accountability to execute.'
)

add_table_from_data(
    ['Role', 'FTE', 'Accountability', 'Timing'],
    [
        ['Founder / ED', '1.0', 'Vision, artist relationships, operations', 'Existing'],
        ['Gallery Manager', '1.0', 'Daily operations, inventory', 'Existing'],
        ['Content (Freelance)', '0.25', 'Photography, social, artist stories', 'Month 1'],
        ['Grant Writer (Contract)', '0.25-0.5', 'Applications, funder relationships', 'Month 2'],
        ['Workshop Instructor', 'Per event', 'Curriculum, talent spotting', 'Month 2-3'],
    ]
)
doc.add_paragraph(
    'Notice what is NOT here: no full-time Content Producer, no Workshop Coordinator, '
    'no expanded ED hours. Those hires are premature. We validate first, then convert '
    'to FTE when evidence supports it.'
)
add_evidence('BCG_PRODUCT_ASSESSMENT.md Section 9 — hiring-before-revenue risk')

doc.add_heading('Governance', level=2)
doc.add_paragraph(
    'A 5-7 member advisory board in Year 1, including 2-3 artist representatives with veto '
    'power on brand collaborations. Full dual-board structure deploys in Year 2-3 when '
    'scale justifies it.'
)
add_evidence('UNDERWRITER_READY_REPORT.md ONE-WAY DOOR #3; condition-4-simplified-governance.md')

doc.add_page_break()

# ================================================================
# SECTION 10: SUCCESS METRICS
# ================================================================
doc.add_heading('10. Success Metrics', level=1)

doc.add_paragraph(
    'We measure outcomes, not outputs. The metrics that matter are whether customers '
    'are getting value and whether the business is sustainable.'
)

doc.add_heading('Customer Value Metrics', level=2)
add_table_from_data(
    ['Metric', 'Year 1 Target', 'How Measured'],
    [
        ['Artist income from Rossi', '$2,500 avg / artist', 'POS + payment records'],
        ['Top-tier artist income', '$15K+ for top 8', 'Individual P&L'],
        ['Workshop satisfaction', '4.5/5', 'Post-workshop survey'],
        ['Repeat collector rate', '20%+', 'Purchase history'],
        ['Artist retention rate', '85%+', 'Annual roster vs prior'],
    ]
)

doc.add_heading('Business Sustainability Metrics', level=2)
add_table_from_data(
    ['Metric', 'Year 1 Target', 'How Measured'],
    [
        ['Total revenue', '$745K', 'Financial statements'],
        ['Online revenue', '$55K', 'Shopify analytics'],
        ['Cash balance (min)', '>$25K at all times', 'Monthly cash report'],
        ['Revenue diversification', 'No stream >50%', 'Revenue breakdown'],
        ['Grant pipeline', '$200K+ applied, $110K received', 'Grant tracking'],
    ]
)

doc.add_heading('90-Day Discovery Targets (Leading Indicators)', level=2)
add_table_from_data(
    ['Metric', '90-Day Target', 'Why It Matters'],
    [
        ['Pilot workshop attendance', '8+ participants', 'Validates discovery funnel'],
        ['First artist drop revenue', '$500+ in 72 hours', 'Validates digital channel'],
        ['Content-to-sale attribution', '1+ documented instance', 'Validates content strategy'],
        ['Artist case studies completed', '3', 'Validates pipeline narrative'],
        ['Brand partnership conversations', '2+', 'Validates collaboration pipeline'],
        ['Partner gallery LOI', '1', 'Validates expansion model'],
    ]
)

doc.add_page_break()

# ================================================================
# SECTION 11: 90-DAY ROADMAP
# ================================================================
doc.add_heading('11. 90-Day Discovery Roadmap', level=1)

doc.add_heading('Phase 1: Foundation (Days 1-30)', level=2)
add_bullet('Pull POS data. Establish revenue baseline. (2 days, $0)')
add_bullet('Hire freelance photographer. Shoot top 50 products. ($2,000)')
add_bullet('Optimize Shopify store. Set up Printful print-on-demand. ($3,500)')
add_bullet('Launch first Artist Drop on Instagram + Shopify. ($0)')
add_bullet('Define brand architecture and guidelines. ($2,000 freelance)')
add_bullet('Survey 20 roster artists on satisfaction and retention. ($0)')
add_bullet('Begin artist case study interviews (3-5 artists). ($0)')
add_bullet('Contact Alberta Street Gallery (Portland) and Gray Loft (Oakland) for '
           'partnership conversations. ($0)')

p = doc.add_paragraph()
run = p.add_run('Phase 1 Gate: ')
run.bold = True
p.add_run('Do we know our baseline? Is the first artist drop signal positive?')

doc.add_heading('Phase 2: Validation (Days 31-60)', level=2)
add_bullet('Run pilot workshop. Document everything. ($500)')
add_bullet('Complete artist case studies. Write them up.')
add_bullet('Execute second and third Artist Drops. Track conversion.')
add_bullet('Pitch 2 brand collaboration partners.')
add_bullet('Submit first grant applications (SFAC if in cycle).')
add_bullet('Build monthly cash flow with actual baseline data.')

p = doc.add_paragraph()
run = p.add_run('Phase 2 Gate: ')
run.bold = True
p.add_run('Did the workshop work? Are online sales tracking?')

doc.add_heading('Phase 3: Package and Present (Days 61-90)', level=2)
add_bullet('Update business plan with all validated data.')
add_bullet('Write financial projections (Section 8) with actual numbers.')
add_bullet('Reconcile all revenue projections across documents.')
add_bullet('Prepare funding presentation with evidence.')
add_bullet('Submit to grant funders and impact investors.')
add_bullet('Present to Rossi board for Year 1 approval.')

doc.add_heading('Total 90-Day Investment', level=2)
add_table_from_data(
    ['Item', 'Cost'],
    [
        ['Freelance brand design', '$2,000'],
        ['Product photography', '$2,000'],
        ['Shopify optimization + print-on-demand', '$3,500'],
        ['Workshop pilot', '$500'],
        ['Miscellaneous', '$2,000'],
        ['TOTAL', '$10,000'],
    ]
)
doc.add_paragraph(
    'This $10,000 investment transforms Rossi from a business plan into a validated, '
    'evidence-backed funding proposal. Expected return: $15-25K in incremental digital '
    'revenue within 90 days, plus the evidence needed to unlock $185K in funding.'
)

doc.add_page_break()

# ================================================================
# SECTION 12: RECOMMENDATION
# ================================================================
doc.add_heading('12. Recommendation to Leadership', level=1)

doc.add_heading('The Honest Assessment', level=2)
doc.add_paragraph(
    'Rossi has built something genuinely valuable. 80+ artists, a Mission District '
    'location with irreplaceable cultural credibility, a 50/50 economic model that is '
    'both ethically differentiated and competitively defensible, and a vision for '
    'systematizing artist career development that, if validated, has no direct competitor.'
)
doc.add_paragraph(
    'Our marketplace research confirms this: across 50+ graffiti artists, 7+ online '
    'platforms, and dozens of physical galleries worldwide, no institution does what Rossi '
    'proposes — systematically discover, develop, and launch graffiti artist careers while '
    'keeping artists as economic partners. The closest models (Alberta Street Gallery, '
    'Side Rail Collective) prove the economics but lack the pipeline and scale vision. '
    'The closest brands (KAWS, Shepard Fairey) prove the market but are individual '
    'operations, not platforms.'
)
add_evidence('GRAFFITI_ARTIST_MARKETPLACE_RESEARCH.md — complete ecosystem mapping')

doc.add_paragraph(
    'The anchor comparable validates the model: Creative Growth Art Center has operated '
    'this exact structure for 50 years — nonprofit gallery, 50/50 split, 140 artists, '
    '$1M+ in annual art sales, $3.26M total revenue. Five years of Form 990 data prove '
    'it works at scale. Rossi is the graffiti and street art version of Creative Growth, '
    'serving an underserved artist community in one of America\'s most culturally significant '
    'neighborhoods for street art.'
)
add_evidence('VALIDATED_COMPARABLES_RESEARCH.md — Finding 1: Creative Growth as anchor comparable')

doc.add_paragraph(
    'But the business plan describes a mature, multi-city, 7-stream operation that does '
    'not yet exist. Two of three products have never been tested. The digital strategy is '
    'underdeveloped. The growth projections are venture-scale on nonprofit capital. And '
    'the most basic piece of evidence — current revenue — has not been collected.'
)

doc.add_heading('The Recommendation', level=2)
p = doc.add_paragraph()
run = p.add_run('Conditional Go. ')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph(
    'Invest 90 days and $10,000 in structured product discovery. Validate the three '
    'core products at the SF flagship. Collect the evidence that transforms this plan '
    'from compelling to fundable. Then seek $185K in staged funding tied to milestones.'
)

doc.add_heading('What Must Be True for This to Work', level=2)
add_bullet('Pull actual POS data within 2 weeks. Non-negotiable.', bold_prefix='1. ')
add_bullet('Run 1 pilot workshop within 6 weeks.', bold_prefix='2. ')
add_bullet('Document 3 real artist case studies.', bold_prefix='3. ')
add_bullet('Launch functioning e-commerce with monthly drops.', bold_prefix='4. ')
add_bullet('Resolve the identity question: cultural platform.', bold_prefix='5. ')
add_bullet('Moderate growth projections to 35-45% annually.', bold_prefix='6. ')
add_bullet('Restructure partner economics before expansion.', bold_prefix='7. ')
add_bullet('Contact Alberta Street (Portland), Gray Loft (Oakland), or Side Rail (Seattle) for LOI.', bold_prefix='8. ')

doc.add_heading('What Rossi Should Do Monday Morning', level=2)
doc.add_paragraph(
    'Pull the POS data. Call three of your top artists. Book a photographer. '
    'Plan the first artist drop. Email Alberta Street Gallery in Portland. '
    'Have the identity conversation. None of this costs more than a few hundred '
    'dollars. All of it moves Rossi from a plan to a product.'
)

p = doc.add_paragraph()
run = p.add_run(
    'The gap is not vision. Rossi has exceptional vision. '
    'The gap is not market — we mapped 50+ artists and the demand is real. '
    'The gap is evidence. Close it in 90 days, and this becomes a funded product.'
)
run.bold = True
run.italic = True

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('---')
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'V2 — Prepared using product-craft methodology and Palette System validation. '
    'Evidence sourced from marketplace research (50+ artists), Creative Growth Form 990 (2020-2024), '
    'validated comparables research, underwriter audit, and BCG assessment. February 2026.'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

# Save
output_path = '/home/mical/fde/rossi-mission-project/ROSSI_PRODUCT_STRATEGY.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')
